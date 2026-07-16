from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# --- 标题 ---
title = doc.add_heading('Python 学习：常见错误记录', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('整理自 2026年7月1日 · 类与继承学习过程').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')


def add_mistake(doc, num, title, wrong_code, why, right_code):
    doc.add_heading(f'错误 {num}：{title}', level=1)

    doc.add_heading('❌ 错误写法', level=2)
    p = doc.add_paragraph()
    code_run = p.add_run(wrong_code)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(10)
    code_run.font.color.rgb = RGBColor(200, 50, 50)

    doc.add_heading('✅ 正确写法', level=2)
    p = doc.add_paragraph()
    code_run = p.add_run(right_code)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(10)
    code_run.font.color.rgb = RGBColor(50, 150, 50)

    doc.add_heading('💡 为什么', level=2)
    doc.add_paragraph(why)

    doc.add_paragraph('─' * 50)


# --- 错误1 ---
add_mistake(doc, 1,
    '用 def 定义类',
    'def IceCreamStand(Restaurant):          # ❌ def 是定义函数的\n    ...',
    'def 用于定义函数，class 用于定义类。Python 看到 def 会认为你在写函数，后面的括号和继承语法都会解析失败。\n\n记忆口诀：类用 class，函数用 def，不混用。',
    'class IceCreamStand(Restaurant):        # ✅ class 才是定义类的\n    ...'
)

# --- 错误2 ---
add_mistake(doc, 2,
    'super.__init__ 缺少括号',
    'super.__init__(restaurant_name, cuisine_type)\n# TypeError: descriptor requires super object but received str',
    'super 是一个类（type），super() 才是创建一个绑定好的 super 对象。\n\n直接用 super.__init__(...) 时，Python 把你传的第一个参数（这里是字符串）当成 self，所以报错 "需要 super 对象但收到了 str"。\n\n记住：super 后面一定有 ()，再点方法名。',
    'super().__init__(restaurant_name, cuisine_type)\n# 先调 super() 创建对象，再调 .__init__()\n# 不需要传 self，Python 自动处理'
)

# --- 错误3 ---
add_mistake(doc, 3,
    '*args 用 append 导致嵌套',
    'def add_flavors(self, *args):\n    self.flavors.append(args)\n\n# 调用 add_flavors("香草", "草莓", "西瓜")\n# 结果: [("香草", "草莓", "西瓜")]  ← 一整个元组塞进去了！',
    '*args 本身就是一个元组。append 会把参数当成"一个整体"加入，所以元组整个变成列表的一个元素 → 嵌套结构。\nextend 会把参数"拆开"逐个加入 → 平铺结构。\n\n选谁取决于数据结构：\n  • 想要列表里套列表 → append\n  • 想要平铺展开 → extend',
    'def add_flavors(self, *args):\n    self.flavors.extend(args)\n\n# 调用 add_flavors("香草", "草莓", "西瓜")\n# 结果: ["香草", "草莓", "西瓜"]  ← 逐个平铺加入'
)

# --- 错误4 ---
add_mistake(doc, 4,
    '子类 super().__init__ 传了未定义的变量',
    'class Admin(User):\n    def __init__(self, first_name, last_name, **kwargs):\n        super().__init__(first_name, last_name, privileges, **kwargs)\n        #                                       ^^^^^^^^^^\n        #                                       NameError! privileges 没定义',
    'privileges 既不是参数，也没有在方法体内定义，Python 找不到这个变量名。\n\n解决方法：privileges 应该作为关键字参数传入 **kwargs，让父类自然接收，然后子类用 pop 提取出来。',
    'class Admin(User):\n    def __init__(self, first_name, last_name, **kwargs):\n        super().__init__(first_name, last_name, **kwargs)\n        self.privileges = self.kwargs.pop("privileges", [])'
)

# --- 错误5 ---
add_mistake(doc, 5,
    '给父类传了过多位置参数',
    '# 父类定义\nclass User:\n    def __init__(self, first_name, last_name, **kwargs):\n        # 只接收 2 个位置参数 ↑\n\n# 子类错误调用\nsuper().__init__(first_name, last_name, privileges, **kwargs)\n#                                 ^^^^^^^^^^  第3个位置参数！TypeError',
    '父类 __init__ 签名明确规定只接收 first_name 和 last_name 两个位置参数。\n传第3个位置参数 privileges 会被 Python 拦截，抛出 TypeError。\n\n解决方法：\n  1. 调用 super().__init__ 时只传父类声明过的位置参数\n  2. 额外的参数让它们进入 **kwargs（关键字参数自动归入）',
    '# 正确：额外参数通过关键字传入\nadmin1 = Admin("sam", "froest", privileges=["删帖", "封号"])\n# privileges 自动进入 **kwargs，父类不会报错'
)

# --- 错误6 ---
add_mistake(doc, 6,
    '循环自赋值 (self.x = self.x)',
    'class Admin(User):\n    def __init__(self, ...):\n        super().__init__(...)\n        self.privileges = self.privileges\n        # ↑ 右边的 self.privileges 从哪来？\n        # AttributeError! 属性还不存在',
    'self.privileges = self.privileges 是死循环逻辑：\n  • 右边 self.privileges 需要从某个地方取值\n  • 但这个属性还没被创建\n  • Python 抛出 AttributeError\n\n自赋值只在一种情况下有意义：x = x + 1（先读旧值再写回）。\n单纯 self.x = self.x 永远没用。',
    'class Admin(User):\n    def __init__(self, ...):\n        super().__init__(...)\n        self.privileges = self.kwargs.pop("privileges", [])\n        # ↑ 明确数据来源：从 kwargs 字典里取，没有就给 []'
)

# --- 错误7 ---
add_mistake(doc, 7,
    'dict.pop() 没给默认值导致 KeyError',
    'self.privileges = self.kwargs.pop("privileges")\n# 如果 "privileges" 不在字典里 → KeyError',
    'dict.pop(key) 当 key 不存在时会抛出 KeyError，程序崩溃。\ndict.pop(key, default) 当 key 不存在时返回默认值，不报错。\n\n在这个场景里，[]（空列表）是合适的默认值：\n  • 管理员没有权限 → 返回空权限列表\n  • 后续代码可以安全地遍历或 append 这个列表\n\n常见默认值：\n  • 列表 → []\n  • 字典 → {}\n  • 数字 → 0\n  • 字符串 → "" 或 None',
    'self.privileges = self.kwargs.pop("privileges", [])\n# 如果 "privileges" 不在字典里 → 返回 [] 保底'
)

# --- 错误8 ---
add_mistake(doc, 8,
    'for 循环边遍历边修改列表',
    '# ❌ 边遍历边删除，会跳过元素！\nnums = [1, 2, 3, 4, 5]\nfor num in nums:\n    if num % 2 == 0:\n        nums.remove(num)\n# 期望: [1, 3, 5]\n# 实际: [1, 3, 5] ← 这次碰巧对了\n# 但如果是 [1, 2, 4, 3, 5]，2被删后4会补到2的位置，循环跳过4！',
    '遍历时删除元素，列表会实时缩短，索引前移，导致下一个元素被跳过。\n\n遍历时删除的后果：\n  • 列表长度一直在变\n  • 当前索引指向的元素已经被换成了下一个\n  • 循环指针又往前一步 → 跳过了一个元素\n\n两种安全做法：\n  1. 遍历副本：for num in nums[:]:  ← 切片创建副本\n  2. 列表推导式：nums = [n for n in nums if n % 2 != 0]\n\n更通用的规则：遍历时不要修改同一个列表的结构。',
    '# ✅ 方法1：遍历副本（切片）\nnums = [1, 2, 4, 3, 5]\nfor num in nums[:]:        # nums[:] 是独立副本\n    if num % 2 == 0:\n        nums.remove(num)     # 在原始列表上删除\nprint(nums)  # [1, 3, 5] ✅\n\n# ✅ 方法2：列表推导式（更 Python 风格）\nnums = [1, 2, 4, 3, 5]\nnums = [n for n in nums if n % 2 != 0]\nprint(nums)  # [1, 3, 5] ✅'
)

# --- 错误9 ---
add_mistake(doc, 9,
    'f-string 内部引号与外层冲突',
    '# ❌ 双引号套双引号 → SyntaxError\nperson = {"name": "张三"}\nprint(f"{person["name"]}")   # Python 3.11 及以前报错！',
    'Python 3.11 及以前版本中，f-string 内部的引号不能和外层相同。\n\n外层用双引号 → 内部字典索引必须用单引号\n外层用单引号 → 内部字典索引必须用双引号\n\nPython 3.12 起已修复这个问题，但公司/服务器环境可能还是老版本，养成习惯总没错。',
    '# ✅ 外层双引号 + 内层单引号\nperson = {"name": "张三"}\nprint(f"{person[\'name\']}")   # ✅ 引号类型不同\n\n# ✅ 或者外层单引号 + 内层双引号\nprint(f\'{person["name"]}\')   # ✅ 也可以'
)

# --- 附录：方法对比 ---
doc.add_heading('附录：列表添加方法对比', level=1)

table = doc.add_table(rows=7, cols=4, style='Light Grid Accent 1')
headers = ['方法', '加了几个元素', '改原列表？', '说明']
for i, text in enumerate(headers):
    table.rows[0].cells[i].text = text

data = [
    ['append(x)', '1 个（x 本身）', '✅ 原地改', '把参数当整体加进去'],
    ['extend(seq)', 'N 个（拆开）', '✅ 原地改', '把可迭代对象逐个加进去'],
    ['insert(i, x)', '1 个', '✅ 原地改', '在指定索引位置插入'],
    ['a + b', '—', '❌ 返回新列表', '合并两个列表，原列表不变'],
    ['a += b', 'N 个', '✅ 原地改', '等价于 extend'],
    ['a[i:i] = seq', 'N 个', '✅ 原地改', '切片赋值，在空区间插入'],
]
for row_idx, row_data in enumerate(data, start=1):
    for col_idx, text in enumerate(row_data):
        table.rows[row_idx].cells[col_idx].text = text

# --- 附录：super() 关键点 ---
doc.add_heading('附录：super() 关键点', level=1)

doc.add_paragraph(
    '1. super 是类，super() 是对象 — 必须先调 super() 创建对象再点方法\n'
    '2. 不要传 self — super() 自动绑定当前实例\n'
    '3. 只传父类声明的参数 — 不要多传位置参数\n'
    '4. 语法模板：'
)

p = doc.add_paragraph()
run = p.add_run(
    'class Child(Parent):\n'
    '    def __init__(self, parent_param1, parent_param2, **kwargs):\n'
    '        super().__init__(parent_param1, parent_param2, **kwargs)\n'
    '        self.child_attr = ...'
)
run.font.name = 'Consolas'
run.font.size = Pt(10)

# --- 保存 ---
save_dir = r'c:\Users\s1040\PycharmProjects\PythonProject\after8.6_learning'
save_path = os.path.join(save_dir, 'Python常见错误记录.docx')
doc.save(save_path)
print(f'已保存: {save_path}')
