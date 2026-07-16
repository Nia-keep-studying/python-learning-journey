from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# --- 样式设置 ---
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)

# --- 标题 ---
title = doc.add_heading('碳循环 + 16:8 间歇断食 减脂计划', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('孙家祥 · 2026年7月2日').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('体重 75kg · 身高 174cm · 体脂 ~25% · 训练时间 18:30').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# ============================================================
# 1. 基础数据
# ============================================================
doc.add_heading('一、基础数据', level=1)

table = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [
    ('身高', '174cm'),
    ('体重', '75kg'),
    ('体脂率', '~25%（瘦体重 56kg / 脂肪 19kg）'),
    ('基础代谢 BMR', '~1,700 kcal（Mifflin 公式）'),
    ('维持热量 TDEE', '~2,000 ~ 2,100 kcal（活动系数 1.2，低能量型）'),
    ('训练节奏', '胸 → 背 → 肩 → 休 · 4天一循环 · 18:30 训练'),
]
for i, (label, value) in enumerate(data):
    table.rows[i].cells[0].text = label
    table.rows[i].cells[1].text = value

doc.add_paragraph('')

# ============================================================
# 2. 16:8 进食窗口
# ============================================================
doc.add_heading('二、16:8 进食窗口', level=1)

doc.add_paragraph('进食窗口：12:00 ~ 20:00（中午到晚上八点）').bold = True
doc.add_paragraph('空腹窗口：20:00 ~ 次日 12:00（16 小时，只喝水/黑咖啡/茶）')

doc.add_paragraph(
    '训练 18:30 正好卡在窗口后半段——练前已在窗口内吃过碳水，有力气；'
    '练后 30 分钟内赶上最后一餐，碳水+蛋白质直接往肌肉里送。'
    '20:00 封嘴后不进食，身体整晚处于脂肪燃烧模式。'
)

doc.add_paragraph('')

# ============================================================
# 3. 碳循环总览
# ============================================================
doc.add_heading('三、4 天碳循环总览', level=1)

table2 = doc.add_table(rows=5, cols=6, style='Light Grid Accent 1')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['天', '训练', '碳水等级', '碳水 (g)', '蛋白质 (g)', '热量 (kcal)']
for i, h in enumerate(headers):
    table2.rows[0].cells[i].text = h
    for p in table2.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True

cycle_data = [
    ['Day 1', '胸', '高碳', '220 (2.9g/kg)', '150', '~1,850'],
    ['Day 2', '背', '高碳', '220 (2.9g/kg)', '150', '~1,850'],
    ['Day 3', '肩', '中碳', '130 (1.7g/kg)', '150', '~1,700'],
    ['Day 4', '休息', '低碳', '50 (0.7g/kg)', '160', '~1,600'],
]
for row_idx, row_data in enumerate(cycle_data, start=1):
    for col_idx, text in enumerate(row_data):
        table2.rows[row_idx].cells[col_idx].text = text

doc.add_paragraph('')

# 脂肪表
doc.add_heading('各日脂肪目标', level=2)
table3 = doc.add_table(rows=4, cols=2, style='Light Grid Accent 1')
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
table3.rows[0].cells[0].text = '日类型'
table3.rows[0].cells[1].text = '脂肪 (g)'
for p in table3.rows[0].cells[0].paragraphs:
    for run in p.runs:
        run.bold = True
for p in table3.rows[0].cells[1].paragraphs:
    for run in p.runs:
        run.bold = True

fat_data = [
    ['高碳日（胸/背）', '45'],
    ['中碳日（肩）', '55'],
    ['低碳日（休息）', '75'],
]
for row_idx, (label, value) in enumerate(fat_data, start=1):
    table3.rows[row_idx].cells[0].text = label
    table3.rows[row_idx].cells[1].text = value

doc.add_paragraph('')

# ============================================================
# 4. 每餐参考（16:8 版）
# ============================================================
doc.add_heading('四、每餐参考（16:8 版：12:00 ~ 20:00）', level=1)

# 高碳日
doc.add_heading('高碳日（胸 / 背）', level=2)

high_data = [
    ['第一餐 12:00', '燕麦 70g + 鸡蛋 2个 + 牛奶 250ml\n（碳水 ~55g / 蛋白质 ~30g）'],
    ['练前加餐 16:30', '米饭 150g（熟）+ 鸡胸肉 120g\n（碳水 ~45g / 蛋白质 ~35g）—— 练前 2h 吃，有力气'],
    ['练后晚餐 19:30', '米饭 150g（熟）+ 鱼/虾 150g + 蔬菜\n（碳水 ~45g / 蛋白质 ~35g）—— 练后 30min 内吃'],
    ['全天补足', '碳水不够的部分用红薯/全麦面包补；蛋白不够用蛋白粉\n20:00 封嘴，只喝水'],
]
table4 = doc.add_table(rows=len(high_data), cols=2, style='Light Grid Accent 1')
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (meal, content) in enumerate(high_data):
    table4.rows[i].cells[0].text = meal
    table4.rows[i].cells[1].text = content

doc.add_paragraph('')

# 中碳日
doc.add_heading('中碳日（肩）', level=2)

med_data = [
    ['第一餐 12:00', '全麦面包 2片 + 鸡蛋 2个 + 牛奶 250ml\n（碳水 ~35g / 蛋白质 ~30g）'],
    ['练前加餐 16:30', '米饭 100g（熟）+ 瘦牛肉 120g\n（碳水 ~30g / 蛋白质 ~30g）'],
    ['练后晚餐 19:30', '土豆 150g + 鸡腿（去皮）150g + 蔬菜\n（碳水 ~35g / 蛋白质 ~40g）'],
    ['全天补足', '碳水还不够用红薯补到 130g；蛋白粉补蛋白缺口'],
]
table5 = doc.add_table(rows=len(med_data), cols=2, style='Light Grid Accent 1')
table5.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (meal, content) in enumerate(med_data):
    table5.rows[i].cells[0].text = meal
    table5.rows[i].cells[1].text = content

doc.add_paragraph('')

# 低碳日
doc.add_heading('低碳日（休息 —— 重点：双重燃脂）', level=2)

low_data = [
    ['第一餐 12:00', '鸡蛋 3个 + 牛油果半个 + 蔬菜\n（碳水 ~10g / 蛋白质 ~18g / 脂肪 ~20g）'],
    ['第二餐 19:30', '牛肉/鸡胸 200g + 大量蔬菜 + 橄榄油 15ml\n（碳水 ~15g / 蛋白质 ~50g / 脂肪 ~30g）'],
    ['额外蛋白', '晚餐后喝一杯蛋白粉（25g 蛋白），凑到 160g\n别让身体拆自己的肌肉'],
    ['全天碳水', '只有 ~25g 来自蔬菜和鸡蛋，其余靠脂肪供能\n配合 16h 空腹 → 整晚纯烧脂肪'],
]
table6 = doc.add_table(rows=len(low_data), cols=2, style='Light Grid Accent 1')
table6.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (meal, content) in enumerate(low_data):
    table6.rows[i].cells[0].text = meal
    table6.rows[i].cells[1].text = content

doc.add_paragraph('')

# 一天时间表
doc.add_heading('一天时间表总览（每个训练日都一样）', level=2)

schedule_data = [
    ['8:00 ~ 12:00', '空腹期 · 只喝水 / 黑咖啡 / 茶（无糖无奶）'],
    ['12:00', '第一餐 —— 破开空腹，身体开始接收今天的热量'],
    ['16:30', '练前加餐 —— 碳水给训练供能，蛋白质防肌肉分解'],
    ['18:30', '训练 💪'],
    ['19:30', '练后晚餐 —— 碳水+蛋白质往肌肉里送，修复开始'],
    ['20:00', '封嘴 —— 只喝水，进入 16 小时燃脂窗口'],
]
table_schedule = doc.add_table(rows=len(schedule_data), cols=2, style='Light Grid Accent 1')
table_schedule.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (time, detail) in enumerate(schedule_data):
    table_schedule.rows[i].cells[0].text = time
    table_schedule.rows[i].cells[1].text = detail

doc.add_paragraph('')

# ============================================================
# 5. 三条铁律
# ============================================================
doc.add_heading('五、三条铁律', level=1)

rules = [
    (
        '蛋白质每天雷打不动 150g+',
        '低碳日可以更多（160g），防止掉肌肉。你目前每天只有 60-70g，先往 100g 冲，适应一周再加。\n'
        '加速补蛋白：蛋白粉一勺 = 25g ❘ 鸡胸肉 100g = 31g ❘ 鸡蛋 1个 = 6g ❘ 牛奶 250ml = 8g ❘ 豆腐 200g = 16g\n'
        '低碳日只有两顿饭 → 肉至少 200g，不够用蛋白粉凑。'
    ),
    (
        '碳水来源选粗的',
        '燕麦、红薯、糙米、全麦面包——别用精制糖填指标。粗粮升糖慢、饱腹感强。\n'
        '低碳日碳水天生只有 ~25g，来自蔬菜和鸡蛋，不需要专门吃碳水食物。'
    ),
    (
        '训练日碳水堆在训练前后',
        '练前 2h 吃碳水（16:30 加餐）→ 训练有力量\n'
        '练后 30min 吃碳水+蛋白质（19:30 晚餐）→ 肌肉修复效率最高\n'
        '非训练窗口的碳水要往前挪，别堆在空腹期之前。'
    ),
]

for i, (title_text, body) in enumerate(rules, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'铁律 {i}：{title_text}')
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph(body)
    doc.add_paragraph('')

# ============================================================
# 6. 体脂高的优势
# ============================================================
doc.add_heading('六、体脂 25% 不是坏事', level=1)

doc.add_paragraph(
    '体脂 25% 在减脂期其实是优势：身体有充足的脂肪储备，低碳日不容易真的"饿"，更多是心理上想吃。'
    '高体脂的人新手期减脂效果特别明显，头两周可能掉 1~2kg（大部分是水，但镜子已经能看出变化）。'
)
doc.add_paragraph(
    '你是低能量型 + 高体脂 → 身体更擅长储存而不是消耗。'
    '这也是为什么碳循环比匀速减脂更适合你 —— 训练日用碳水把代谢拉上去，休息日用低碳 + 空腹把脂肪烧下去。'
)

doc.add_paragraph('')

# ============================================================
# 7. 每日记录表
# ============================================================
doc.add_heading('七、每日记录表', level=1)
doc.add_paragraph('每天填一行，观察趋势。复制此表追加到下一页。')

table7 = doc.add_table(rows=2, cols=7, style='Light Grid Accent 1')
table7.alignment = WD_TABLE_ALIGNMENT.CENTER
record_headers = ['日期', '天', '训练', '体重 (kg)', '蛋白质 (g)', '碳水 (g)', '备注']
for i, h in enumerate(record_headers):
    table7.rows[0].cells[i].text = h
    for p in table7.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True
# 示例行
example = ['7/7（一）', 'Day 1', '胸', '75.0', '', '', '第一轮开始！']
for i, v in enumerate(example):
    table7.rows[1].cells[i].text = v

doc.add_paragraph('')

# ============================================================
# 8. 调整规则
# ============================================================
doc.add_heading('八、调整规则', level=1)

adjustments = [
    '每周日早上空腹称一次体重，每周拍一张正面照（镜子不撒谎）。',
    '掉太慢（<0.3kg/周）→ 低碳日再加一天（变成胸-背-肩-休-休，中碳去掉），或高碳日碳水减到 180g。',
    '掉太快（>1kg/周）→ 多是水和肌肉！高碳日碳水拉到 250g，或中碳日拉到 160g。',
    '训练力量明显下降 → 练前碳水多加 30g，或把进食窗口往前挪 1 小时。',
    '低碳日两顿饭吃不饱 → 蔬菜无限量，橄榄油多放 5ml，或者蛋白粉分两次喝。',
    '16:8 如果上午实在太饿 → 窗口改成 10:00-18:00（跳过晚餐），效果一样。',
]
for adj in adjustments:
    doc.add_paragraph(adj, style='List Bullet')

# --- 保存到桌面 ---
desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
save_path = os.path.join(desktop, '碳循环减脂计划.docx')
doc.save(save_path)
print(f'Done: {save_path}')
